import jieba
from gensim.models import word2vec
import os,re,csv
import docx
from GTS import graces
import pandas as pd
import pdfminer

from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfdevice import PDFDevice
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document



def get_cut_word_dict(patients_Context_path,medical_dict):
    fout_regex = open(medical_dict, 'a', encoding='utf8', newline='')
    lst1 = []
    for root, dirs, files in os.walk(patients_Context_path):
        if files:
            print(root.split("\\")[-1])
            patient_root = os.path.join(root, '首程.docx')
            for i in docx.Document(patient_root).paragraphs:
                p1 = re.compile(r'\d+[次度年岁月]').findall(i.text)
                p2 = re.compile(r'([a-zA-Z0-9+]+[\.^]*[A-Za-z0-9%(℃)]+(?![次度年岁分钟]))').findall(i.text)
                p_merge = p1 + p2
                # print(p_merge)
                if len(p_merge):
                    for word in p_merge:
                        lst1.append(word)
    lst2 = list(set(lst1))
    for i in lst2:
        fout_regex.write('\n' + i)
    fout_regex.close()



patients_Context_path = './data/patients_data1/'
jieba.load_userdict('./data/medical_ner_dict.txt')
patient_context_sentens = []
stopwords = [' ','、','：','；']
# for root, dirs ,files in os.walk(patients_Context_path):
#     if files:
#         print(root.split("\\")[-1])
#         patient_root = os.path.join(root,'首程.docx')
#         for i in docx.Document(patient_root).paragraphs:
#             sentence_list = []
#             for word in jieba.lcut(i.text.replace("，","").replace(',',"").replace("。","").replace('"',"")):
#                 if word not in stopwords:
#                     sentence_list.append(word)
#             patient_context_sentens.append(sentence_list)
#
# model = word2vec.Word2Vec(patient_context_sentens, sg=0, epochs=8,vector_size=512,  window=5, negative=3, sample=0.001, hs=1, workers=16,min_count=1)
# model.save('./data/models/word2vec.model')



# df = pd.DataFrame(columns =['patient_name','first_trip','pdf'])
# patients_Context_path = './data/patients_data'
# for root, dirs ,files in os.walk(patients_Context_path):
#     if files:
#         ### pdf
#         print(root.split("\\")[-1])
#         df_len = len(df)
#         df.loc[df_len,'patient_name'] = root.split("\\")[-1]
#         patient_root = os.path.join(root,'document.pdf')
#         with open(patient_root,"rb") as pdf:
#             parser = PDFParser(pdf)
#             # 创建一个PDF文档
#             doc = PDFDocument()
#             # 分析器和文档相互连接
#             parser.set_document(doc)
#             doc.set_parser(parser)
#             # 提供初始化密码，没有默认为空
#             doc.initialize()
#             # 检查文档是否可以转成TXT，如果不可以就忽略
#             if not doc.is_extractable:
#                 raise PDFTextExtractionNotAllowed
#             else:
#                 # 创建PDF资源管理器，来管理共享资源
#                 rsrcmagr = PDFResourceManager()
#                 # 创建一个PDF设备对象
#                 laparams = LAParams()
#                 # 将资源管理器和设备对象聚合
#                 device = PDFPageAggregator(rsrcmagr, laparams=laparams)
#                 # 创建一个PDF解释器对象
#                 interpreter = PDFPageInterpreter(rsrcmagr, device)
#             flag_page = True
#             flag_sentence = False
#             patient_content = ''
#             for page in doc.get_pages():
#                 if flag_page:
#                     interpreter.process_page(page)
#                     # 接收该页面的LTPage对象
#                     layout = device.get_result()
#                     # 这里的layout是一个LTPage对象 里面存放着page解析出来的各种对象
#                     # 一般包括LTTextBox，LTFigure，LTImage，LTTextBoxHorizontal等等一些对象
#                     # 想要获取文本就得获取对象的text属性
#                     for x in layout:
#                         if (isinstance(x, LTTextBoxHorizontal)):
#                             #print(x.get_text().split('\n')[0])
#                             if flag_sentence:
#                                 reg = re.compile(r'.*术后第*\d*[二|2]*天')
#                                 if re.match(reg,x.get_text().split('\n')[0]):
#                                     # print("finish!")
#                                     flag_sentence = False
#                                     break
#                                 else:
#                                     patient_content = patient_content + x.get_text().split('\n')[0]
#                             else:
#                                 reg = re.compile(r'.*术后第*[1|一]天')
#                                 if re.match(reg,x.get_text().split('\n')[0]):
#                                     # print("match!!")
#                                     flag_sentence = True
#                                     patient_content = x.get_text().split('\n')[0]
#                 else:
#                     break
#         # print(patient_content.rsplit('。',1)[0].strip())
#         df.loc[df_len, 'pdf'] = patient_content.rsplit('。',1)[0].strip()
#         ### docx
#         patient_docx_root = os.path.join(root, '首程.docx')
#         patient_docx = ''.join([sentence.text for sentence in docx.Document(patient_docx_root).paragraphs])
#         df.loc[df_len, 'first_trip'] = patient_docx
#
# df.to_excel('./data/NLP_patient_context.xlsx')

df = pd.DataFrame(columns =['patient_name','first_trip','pdf'])
patients_Context_path = './data/patients_data2'
for root, dirs ,files in os.walk(patients_Context_path):
    if files:
        name = root.split("\\")[-1]
        patient_root = os.path.join(root, 'document.docx')
        if os.path.exists(patient_root):
            Contexts = docx.Document(patient_root).paragraphs
            print(name)
            df_len = len(df)
            df.loc[df_len,'patient_name'] = root.split("\\")[-1]
            is_rematch = False

            for sentence in Contexts:
                if is_rematch:
                    reg = re.compile(r'.*诊*断依据*')
                    if re.match(reg,sentence.text.split('\n')[0]):
                        # print("finish!")
                        flag_sentence = False
                        print(sentence.text)
                        break
                    else:
                        patient_content = patient_content + sentence.text.split('\n')[0]

                else:
                    reg = re.compile(r'.*患者.*')
                    if re.match(reg, sentence.text.split('\n')[0]):
                        is_rematch = True
                        patient_content = sentence.text.split('\n')[0]

            df.loc[df_len,'first_trip'] = patient_content
            is_rematch = False
            for sentence in Contexts:
                if is_rematch:
                    reg = re.compile(r'.*术后第*\d*[二|2]*天')
                    if re.match(reg,sentence.text.split('\n')[0]):
                        # print("finish!")
                        flag_sentence = False
                        print(sentence.text)
                        break
                    else:
                        patient_content = patient_content + sentence.text.split('\n')[0]

                else:
                    reg = re.compile(r'.*术后第*[1|一]天*')
                    if re.match(reg, sentence.text.split('\n')[0]):
                        is_rematch = True
                        patient_content = sentence.text.split('\n')[0]
            df.loc[df_len, 'pdf'] = patient_content
    df_first = pd.read_excel('./data/NLP_patient_context.xlsx',index_col=0)
    # df = pd.read_excel('./data/check.xlsx',index_col=0)
    df_concat = pd.concat([df_first,df])
    df_concat.to_excel('./data/NLP_patient_context_concat.xlsx')
    # df.to_excel('./data/check.xlsx')




























