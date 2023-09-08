import jieba
from gensim.models import word2vec
import os,re,csv
import docx
# from GTS import graces
import pandas as pd
import pdfminer
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfdevice import PDFDevice
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal, LAParams
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document
if __name__ == '__main__':
    df = pd.DataFrame(columns=['patient_name', 'dis_day_pdf'])
    patients_Context_path = './data/patients_data1'
    for root, dirs ,files in os.walk(patients_Context_path):
        if files:
            name = root.split("\\")[-1]
            patient_root = os.path.join(root, 'document.docx')
            if os.path.exists(patient_root):
                Contexts = docx.Document(patient_root).paragraphs
                is_rematch = False
                # print()
                signature_list = []
                patient_content = ''
                for index in range(len(Contexts) - 1, 0, -1):
                    sentence = Contexts[index]
                    if is_rematch:
                        reg = re.compile(r'.*今.*查房.*')
                        if re.match(reg, sentence.text.split('\n')[0]):
                            patient_content = sentence.text.split('\n')[0] + patient_content
                            break
                        else:
                            reg_signature = re.compile(r'.*医师签名.*')
                            if re.match(reg_signature, sentence.text.split('\n')[0]):
                                pass
                            else:
                                patient_content = sentence.text.split('\n')[0].strip() + patient_content
                    else:
                        reg = re.compile(r'.*今.*查房.*')
                        if re.match(reg, sentence.text.split('\n')[0]):
                            is_rematch = True
                df_len = len(df)
                df.loc[df_len, 'patient_name'] = name
                df.loc[df_len, 'dis_day_pdf'] = patient_content


    patients_Context_path = './data/patients_data2'
    for root, dirs, files in os.walk(patients_Context_path):
        if files:
            name = root.split("\\")[-1]
            patient_root = os.path.join(root, 'document.docx')
            if os.path.exists(patient_root):
                Contexts = docx.Document(patient_root).paragraphs
                is_rematch = False
                # print()
                signature_list = []
                patient_content = ''
                for index in range(len(Contexts) - 1, 0, -1):
                    sentence = Contexts[index]
                    if is_rematch:
                        reg = re.compile(r'.*今.*查房.*')
                        if re.match(reg, sentence.text.split('\n')[0]):
                            patient_content = sentence.text.split('\n')[0] + patient_content
                            break
                        else:
                            reg_signature = re.compile(r'.*医师签名.*')
                            if re.match(reg_signature, sentence.text.split('\n')[0]):
                                pass
                            else:
                                 patient_content = sentence.text.split('\n')[0].strip() + patient_content
                    else:
                        reg = re.compile(r'.*今.*查房.*')
                        if re.match(reg, sentence.text.split('\n')[0]):
                            is_rematch = True
                df_len = len(df)
                df.loc[df_len, 'patient_name'] = name
                df.loc[df_len, 'dis_day_pdf'] = patient_content
                # for sentence in Contexts:
                #     reg = re.compile(r'.*医师签名：.*')
                #     if re.match(reg, sentence.text.split('\n')[0]):
                #         context_rematch = sentence.text.split('\n')[0].split('：')[1].strip().replace('医师签名','')
                #         for i in context_rematch.split('/'):
                #             signature.append(i)

    df.to_excel('./data/NLP_patient_dis_day.xlsx')